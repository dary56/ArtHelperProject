import os
import tempfile
import pypandoc
from django.core.files.base import ContentFile
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm


def _insert_empty_paragraph_before(element):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    p.append(r)
    element.addprevious(p)

def _insert_empty_paragraph_after(element):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    p.append(r)
    element.addnext(p)

def _get_element_text(element):
    return ''.join(t.text for t in element.iter(qn('w:t')) if t.text)

def _is_image_table_element(tbl_element):
    rows = tbl_element.findall(qn('w:tr'))
    if len(rows) != 1:
        return False
    cells = rows[0].findall(qn('w:tc'))
    if len(cells) != 1:
        return False
    return 'Рисунок' in _get_element_text(cells[0])

def _insert_empty_lines(doc):
    for element in list(doc.element.body):
        tag = element.tag
        if tag == qn('w:tbl'):
            if _is_image_table_element(element):
                _insert_empty_paragraph_before(element)   # перед рисунком
                _insert_empty_paragraph_after(element)      # после подписи рисунка
            else:
                _insert_empty_paragraph_after(element)      # после обычной таблицы
        elif tag == qn('w:p'):
            if _get_element_text(element).startswith('Таблица '):
                _insert_empty_paragraph_before(element)   # перед подписью таблицы


def _remove_table_borders(table):
    """Убирает все видимые границы таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_table_borders(table):
    """Добавляет чёрные сплошные границы таблицы."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _is_image_table(table):
    """Проверяет, является ли таблица обёрткой Pandoc для figure (рисунок)."""
    if len(table.rows) != 1:
        return False
    if len(table.rows[0].cells) != 1:
        return False
    cell_text = table.rows[0].cells[0].text.strip()
    return 'Рисунок' in cell_text


def _format_image_table(table):
    """Убирает границы, центрирует содержимое таблицы-рисунка."""
    _remove_table_borders(table)
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    
    # Убираем фиксированную ширину
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    
    # Центрируем таблицу на странице
    existing_jc = tblPr.find(qn('w:jc'))
    if existing_jc is not None:
        tblPr.remove(existing_jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    tblPr.append(jc)
    
    # Форматируем ячейку
    cell = table.rows[0].cells[0]
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(10)


def _format_normal_table(table):
    """Обычная таблица: границы, 100% ширина, шрифт 10, выравнивание по ширине."""
    _set_table_borders(table)
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)
    
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), '5000')
    tblPr.append(tblW)
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0)
                paragraph.paragraph_format.right_indent = Cm(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def _format_captions(doc):
    """Форматирует подписи рисунков и таблиц: шрифт 10, по центру, без отступа."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text.startswith('Рисунок '):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.italic = True
                run.font.bold = False

        elif text.startswith('Таблица '):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0)
            paragraph.paragraph_format.right_indent = Cm(0)
            for run in paragraph.runs:
                run.font.size = Pt(10)
                run.font.bold = False
                run.font.italic = False


def export_article_docx(article):
    from .html_builder import build_article_html

    html = build_article_html(article)
    template_path = article.metadata.journal.get_active_template().file.path

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name

    try:
        pypandoc.convert_text(
            html,
            'docx',
            format='html',
            outputfile=tmp_path,
            extra_args=[f'--reference-doc={template_path}']
        )

        # Пост-обработка
        doc = Document(tmp_path)
        for table in doc.tables:
            if _is_image_table(table):
                _format_image_table(table)
            else:
                _format_normal_table(table)
        _format_captions(doc)
        _insert_empty_lines(doc)
        doc.save(tmp_path)

        with open(tmp_path, 'rb') as f:
            return ContentFile(f.read())

    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)